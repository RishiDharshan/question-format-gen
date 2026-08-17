"""
run_jkfaa.py
------------
Runner for the JKFAA (JK Accounts/Finance Full) question generation.

Repetition-avoidance strategy:
  - Instead of generating all topics in one shot, the chapters list is split
    into batches of `topics_per_batch` chapters (default 3).
  - Each batch is its own generation round so the LLM works over a small,
    focused set of topics, dramatically reducing cross-topic repetition.
  - The deduplication layer in doc_maker still filters any accidental
    duplicates across batches at the DB level.
"""

import asyncio
import json
import os
import time
import copy
from pathlib import Path

from doc_maker import generate_and_make_mock, make_docs, number_questions
from generation_engine import get_individual_jobs, main_worker, load_json
import docx
import re


CONFIG_PATH = "content_meta/jkfaa.json"


def clean_xml(s: str) -> str:
    if s is None:
        return ""
    if not isinstance(s, str):
        s = str(s)

    def ok(cp: int) -> bool:
        return (
            cp == 0x9 or cp == 0xA or cp == 0xD
            or (0x20 <= cp <= 0xD7FF)
            or (0xE000 <= cp <= 0xFFFD)
            or (0x10000 <= cp <= 0x10FFFF)
        )

    return "".join(ch for ch in s if ok(ord(ch)))


def _add_content_simple(doc, text: str):
    """Add text content to the docx document paragraph by paragraph."""
    for line in text.split("\n"):
        doc.add_paragraph(line)


def chunk_list(lst, chunk_size):
    """Yield successive chunks of `chunk_size` from `lst`."""
    for i in range(0, len(lst), chunk_size):
        yield lst[i : i + chunk_size]


async def run():
    # ── Load config ────────────────────────────────────────────────────────
    with open(CONFIG_PATH, "r") as f:
        data = json.load(f)

    db_file = data.get("db_file")
    outputs_dir = data.get("outputs_dir", "outputs")
    topics_per_batch = data.get("topics_per_batch", 3)

    print(f"JKFAA Generation starting — config: {CONFIG_PATH}")
    print(f"  topics_per_batch = {topics_per_batch}")
    print(f"  db_file          = {db_file}")

    # ── Extract chapter list & question-type distribution ──────────────────
    subjects = data["subjects"]
    original_subject = subjects[0]
    all_chapters = original_subject["chapters"]
    types_of_questions = original_subject["types_of_questions"]

    total_distribution = sum(t.get("distribution", 0) for t in types_of_questions)
    total_chapters = len(all_chapters)

    print(f"  total chapters   = {total_chapters}")
    print(f"  total questions  = {total_distribution}")
    print()

    # ── Split chapters into batches of `topics_per_batch` ──────────────────
    chapter_batches = list(chunk_list(all_chapters, topics_per_batch))
    num_batches = len(chapter_batches)

    # Proportionally divide distributions across batches
    # Each batch gets a fair share; remainder goes to first batch
    base_distributions = []
    for q_type in types_of_questions:
        dist = q_type.get("distribution", 0)
        per_batch = dist // num_batches
        remainder = dist % num_batches
        base_distributions.append((per_batch, remainder))

    all_questions = []

    for batch_idx, chapter_batch in enumerate(chapter_batches):
        batch_num = batch_idx + 1
        print(f"── Batch {batch_num}/{num_batches}: {len(chapter_batch)} chapters ──")
        for ch in chapter_batch:
            print(f"   • {ch}")

        # Build a modified config for this batch only
        batch_types = []
        for i, q_type in enumerate(types_of_questions):
            per_batch, remainder = base_distributions[i]
            # Give remainder questions to the first batch
            batch_dist = per_batch + (remainder if batch_idx == 0 else 0)
            if batch_dist == 0:
                # Ensure at least 1 question per type in non-zero batches
                batch_dist = 1 if per_batch == 0 and batch_idx == 0 else per_batch
            batch_types.append({**q_type, "distribution": batch_dist})

        batch_subject = {
            **original_subject,
            "chapters": chapter_batch,
            "types_of_questions": batch_types,
        }

        batch_data = {
            **data,
            "subjects": [batch_subject],
        }

        # Write a temp config for this batch
        temp_config_path = f"content_meta/jkfaa_batch_{batch_idx}.json"
        with open(temp_config_path, "w") as f:
            json.dump(batch_data, f, indent=2)

        try:
            # Generate using existing pipeline
            subject_prompt_map = await get_individual_jobs(data=batch_data)
            if subject_prompt_map:
                result_pool = await main_worker(
                    prompt_maps=subject_prompt_map, db_file=db_file
                )
                for result in result_pool:
                    batch_questions = await make_docs(raw_response=result, db_file=db_file)
                    all_questions.extend(batch_questions)
        finally:
            # Clean up temp config
            if os.path.exists(temp_config_path):
                os.remove(temp_config_path)

        print(f"   Batch {batch_num} done — running total: {len(all_questions)} questions\n")

    # ── Number & save to DOCX ──────────────────────────────────────────────
    all_questions = number_questions(all_questions)

    output_filename = f"jkfaa_output_{int(time.time())}.docx"
    docx_save_path = Path(outputs_dir) / output_filename

    doc = docx.Document()
    formatted = "\n\n".join(all_questions)
    pattern = r"(\*\*|__)(.*?)\1"
    clean_text = re.sub(pattern, r"\2", formatted)

    doc.add_paragraph(clean_xml(clean_text))
    doc.save(path_or_stream=docx_save_path)

    print(f"\n✓ Successfully generated {len(all_questions)} questions.")
    print(f"✓ Saved to: {docx_save_path}")


if __name__ == "__main__":
    asyncio.run(run())
