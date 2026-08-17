import re, asyncio, random
from question_map import get_question_bank
from generation_engine import get_questions
from pathlib import Path
import docx
from openai import AsyncOpenAI
from db_ope import (
    split_question_blocks,
    dedup_blocks_against_db,
    extract_question_line,
    insert_questions,
    insert_embedding,
    insert_concept,
    hash_question,
    SEMANTIC_SIMILARITY_THRESHOLD,
)
from similarity import is_too_similar, get_embedding
from output_schema import ConceptOutput
from generator import agenerate
from dotenv import load_dotenv

load_dotenv()

_EMBED_CLIENT = AsyncOpenAI()

# Rejection reason labels for structured logging
REASON_EXACT_DUPLICATE    = "EXACT_DUPLICATE"
REASON_SEMANTIC_DUPLICATE = "SEMANTIC_DUPLICATE"
REASON_ACCEPTED           = "ACCEPTED"

async def _extract_and_store_concept(db_file: str, qhash: str, qtext: str) -> None:
    """Best-effort: extract the core concept tested by a question and store it."""
    try:
        concept_result = await agenerate(
            client=_EMBED_CLIENT,
            user_prompt=(
                f"Extract the single core concept being tested in this question:\n\n{qtext}"
            ),
            response_model=ConceptOutput,
        )
        if hasattr(concept_result, "concept") and concept_result.concept:
            insert_concept(db_file=db_file, qhash=qhash, concept=concept_result.concept)
    except Exception:
        pass  # concept extraction is best-effort; never fail the pipeline


async def make_docs(raw_response, db_file):
    blocks = split_question_blocks(raw_response)
    if not blocks:
        # No recognisable question blocks; pass through as-is
        final_cleaned = raw_response
        question_lines = final_cleaned.splitlines()
        regex = r'^\(\d+\)\.\s'
        question_map = get_question_bank(text_lines=question_lines, re_exprs=regex)
        return list(question_map.values())

    file_lock = asyncio.Lock()

    # ── Stage 1: Exact-hash deduplication ────────────────────────────────────
    async with file_lock:
        unique_blocks, dup_blocks = dedup_blocks_against_db(blocks, db_file=db_file)

    for b in dup_blocks:
        qline = extract_question_line(b)
        print(f"  [REJECTED:{REASON_EXACT_DUPLICATE}] {qline[:80]!r}")

    # ── Stage 2: Semantic similarity check ───────────────────────────────────
    semantic_unique_blocks = []
    semantic_dup_blocks = []

    for block in unique_blocks:
        qline = extract_question_line(block)
        try:
            too_similar, score, embedding = await is_too_similar(
                client=_EMBED_CLIENT,
                new_question_text=qline,
                db_file=db_file,
                threshold=SEMANTIC_SIMILARITY_THRESHOLD,
            )
        except Exception as e:
            # If the embedding call fails, let the question through (fail-open)
            print(f"  [WARN] Semantic check failed for {qline[:60]!r}: {e}")
            too_similar, score, embedding = False, 0.0, []

        if too_similar:
            semantic_dup_blocks.append(block)
            print(
                f"  [REJECTED:{REASON_SEMANTIC_DUPLICATE}] score={score:.3f} "
                f"threshold={SEMANTIC_SIMILARITY_THRESHOLD} | {qline[:70]!r}"
            )
        else:
            semantic_unique_blocks.append((block, embedding))
            print(f"  [{REASON_ACCEPTED}] score={score:.3f} | {qline[:70]!r}")

    # ── Stage 3: Store accepted questions, embeddings, and concepts ───────────
    concept_tasks = []
    stored_qlines = []

    if semantic_unique_blocks:
        async with file_lock:
            accepted_qlines = [extract_question_line(b) for b, _ in semantic_unique_blocks]
            insert_questions(accepted_qlines, db_file=db_file)
            print(f"  → Stored {len(accepted_qlines)} new question(s) in {db_file}")

        for qline, (block, embedding) in zip(accepted_qlines, semantic_unique_blocks):
            qhash = hash_question(qline)
            # Store embedding (non-blocking, best-effort)
            if embedding:
                try:
                    insert_embedding(db_file=db_file, qhash=qhash, embedding=embedding)
                except Exception:
                    pass
            # Schedule concept extraction (run concurrently after this loop)
            concept_tasks.append(_extract_and_store_concept(db_file=db_file, qhash=qhash, qtext=qline))

        stored_qlines = accepted_qlines

    # Run concept extraction concurrently for all accepted questions
    if concept_tasks:
        await asyncio.gather(*concept_tasks, return_exceptions=True)

    if dup_blocks:
        print(f"  → {len(dup_blocks)} exact duplicate(s) skipped.")
    if semantic_dup_blocks:
        print(f"  → {len(semantic_dup_blocks)} semantic duplicate(s) skipped.")

    # ── Build output ──────────────────────────────────────────────────────────
    final_blocks = [b for b, _ in semantic_unique_blocks]
    final_cleaned = "\n".join(final_blocks)

    question_lines = final_cleaned.splitlines()
    regex = r'^\(\d+\)\.\s'
    question_map = get_question_bank(text_lines=question_lines, re_exprs=regex)
    return list(question_map.values())



def clean_xml(s: str) -> str:
    if s is None:
        return ""
    if not isinstance(s, str):
        s = str(s)

    def ok(cp: int) -> bool:
        return (
            cp == 0x9 or cp == 0xA or cp == 0xD
            or (0x20 <= cp <= 0xD7FF)
            or (0xE000 <= cp <= 0xFFFD)
            or (0x10000 <= cp <= 0x10FFFF)
        )

    return "".join(ch for ch in s if ok(ord(ch)))


def number_questions(questions: list[str]) -> list[str]:
    no_of_questions = len(questions)
    
    for i in range(0, no_of_questions):
        questions[i] = f"{i+1}. "+questions[i].strip()
    
    return questions

def _parse_markdown_table(lines: list[str]) -> list[list[str]]:
    """
    Parse a list of markdown table lines into a list of rows (each row is a list of cell strings).
    Skips separator lines (lines like |---|---|).
    """
    rows = []
    for line in lines:
        stripped = line.strip()
        # Skip separator lines like |---|---| or | --- | --- |
        if re.match(r'^\|[\s\-:]+\|[\s\-:|]*$', stripped):
            continue
        # Split by pipe, strip each cell, remove empty first/last from leading/trailing pipes
        cells = [cell.strip() for cell in stripped.split('|')]
        # Remove empty strings caused by leading/trailing pipes
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        if cells:
            rows.append(cells)
    return rows

def _add_content_with_tables(doc, text: str):
    """
    Parse the full text, detect markdown table blocks, and add them as
    proper Word tables. Non-table text is added as regular paragraphs.
    """
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.enum.table import WD_TABLE_ALIGNMENT

    # Pre-process: split single-line compressed markdown tables into multi-line
    # Detects patterns like "| col1 | col2 | |---|---| | A | B |" on one line
    # and splits at row boundaries (where "| |" or "| \n|" indicates a new row)
    processed_lines = []
    for line in text.split('\n'):
        # Check if this line contains multiple markdown table rows compressed into one
        # Pattern: a line with 3+ pipe characters and a separator pattern like |---|
        if line.count('|') >= 6 and re.search(r'\|[\s\-:]+\|', line):
            # Split at points where a cell ends and a new row starts: "| |" pattern
            # This handles: "| Header1 | Header2 | |---|---| | A | I |"
            parts = re.split(r'\|\s*\|', line)
            if len(parts) > 1:
                # Reconstruct each part as a proper table line with pipes
                reconstructed = []
                for i, part in enumerate(parts):
                    part = part.strip()
                    if not part:
                        continue
                    # Add leading/trailing pipes if missing
                    if not part.startswith('|'):
                        part = '| ' + part
                    if not part.endswith('|'):
                        part = part + ' |'
                    reconstructed.append(part)
                processed_lines.extend(reconstructed)
            else:
                processed_lines.append(line)
        else:
            processed_lines.append(line)
    # A markdown table line starts and ends with a pipe character
    md_table_line_re = re.compile(r'^\s*\|.*\|\s*$')

    segments = []  # list of ('text', str) or ('table', [lines])
    current_text_lines = []
    current_table_lines = []

    for line in processed_lines:
        is_table_line = bool(md_table_line_re.match(line))
        if is_table_line:
            # If we were accumulating text, flush it
            if current_text_lines:
                segments.append(('text', '\n'.join(current_text_lines)))
                current_text_lines = []
            current_table_lines.append(line)
        else:
            # If we were accumulating table lines, flush them
            if current_table_lines:
                segments.append(('table', current_table_lines))
                current_table_lines = []
            current_text_lines.append(line)

    # Flush remaining
    if current_text_lines:
        segments.append(('text', '\n'.join(current_text_lines)))
    if current_table_lines:
        segments.append(('table', current_table_lines))

    for seg_type, seg_data in segments:
        if seg_type == 'text':
            # Add non-empty text as a paragraph
            content = seg_data.strip()
            if content:
                doc.add_paragraph(content)
        elif seg_type == 'table':
            rows = _parse_markdown_table(seg_data)
            if not rows:
                continue
            # Determine number of columns from the first row
            num_cols = max(len(row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Add cell content
            for r_idx, row_cells in enumerate(rows):
                for c_idx, cell_text in enumerate(row_cells):
                    if c_idx < num_cols:
                        cell = table.cell(r_idx, c_idx)
                        cell.text = cell_text

            # Apply borders to all cells for clean table appearance
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is None:
                        from docx.oxml import OxmlElement
                        tcBorders = OxmlElement('w:tcBorders')
                        tcPr.append(tcBorders)
                    for edge in ('top', 'left', 'bottom', 'right'):
                        element = OxmlElement(f'w:{edge}')
                        element.set(qn('w:val'), 'single')
                        element.set(qn('w:sz'), '4')
                        element.set(qn('w:space'), '0')
                        element.set(qn('w:color'), '000000')
                        tcBorders.append(element)

            # Bold the first row (header)
            if table.rows:
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

            doc.add_paragraph('')  # spacer after table

async def generate_and_make_mock(json_path: str, db_file: str, doc_save_path: Path, output_filename: str = None):
    result_pool = await get_questions(json_path=json_path)
    questions = []
    for result in result_pool:
        final_questions = await make_docs(raw_response=result, db_file=db_file)
        questions.extend(final_questions)
        
    questions = number_questions(questions=questions)
        
    if output_filename:
        docx_save_path = doc_save_path / output_filename
    else:
        random_string = ''.join([str(random.randint(0,9)) for _ in range(10)])
        docx_save_path = doc_save_path / f'sample_{random_string}.docx'

    doc = docx.Document()

    formatted_questions = '\n\n'.join(questions)
    pattern = r"(\*\*|__)(.*?)\1"
    clean_text = re.sub(pattern, r"\2", formatted_questions)

    try:
        _add_content_with_tables(doc, clean_xml(clean_text))
        doc.save(path_or_stream=docx_save_path)
    except Exception as e:
        print(f"Error: {e}")





if __name__ == "__main__":
    import json
    import sys
    default_json = '/Users/rishidharshan/projects/question format gen/universal-maybe-main/content_meta/jkssb_full.json'
    json_path = sys.argv[1] if len(sys.argv) > 1 else default_json
    
    with open(json_path, 'r') as f:
        data = json.load(f)
    try:
        db_file = data.get("db_file")
        outputs_dir = data.get("outputs_dir")
        doc_save_path = Path(outputs_dir)
    except Exception as e:
        print("Exception occured: {e}")
        import sys
        sys.exit(1)
    asyncio.run(generate_and_make_mock(json_path=json_path, db_file=db_file, doc_save_path=doc_save_path))



    

